
import docx, pprint


file1 = 'D:\\Andrii Buvailo\\EnamineStore\\PunchOUt\\e-commerce.docx'

doc = docx.Document(file1)
print (len(doc.paragraphs))

print(doc.paragraphs[0].text)


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

pprint.pprint (getText(file1))